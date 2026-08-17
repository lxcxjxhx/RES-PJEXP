#!/usr/bin/env python3
"""Convert resume.html to DOCX with Chinese font support and precise A4 layout."""

import subprocess
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from bs4 import BeautifulSoup


def get_cjk_font():
    """Detect installed CJK font family from fontconfig."""
    try:
        result = subprocess.run(
            ["fc-list", ":lang=zh", "family"],
            capture_output=True, text=True, timeout=10
        )
        families = sorted(set(
            line.strip() for line in result.stdout.strip().split("\n") if line.strip()
        ))
        # Prefer Noto Sans CJK SC
        for f in families:
            if "Noto Sans CJK SC" in f:
                return "Noto Sans CJK SC"
        # Fallback to first available
        if families:
            return families[0]
    except Exception as e:
        print(f"Warning: Could not detect fonts: {e}")
    return "Noto Sans CJK SC"  # Default


FONT = get_cjk_font()
print(f"Using font: {FONT}")

# Colors
DARK = (44, 62, 80)
GRAY = (127, 140, 141)
TEXT = (51, 51, 51)
LINK = (52, 152, 219)
LIGHT_GRAY = (85, 85, 85)
FOOTER_COLOR = (149, 165, 166)


def set_font(run, size=None, bold=False, color=None):
    """Set font properties for a run."""
    run.font.name = FONT
    # Safely set East Asian font - rPr might not exist yet
    from docx.oxml import OxmlElement
    r_elem = run._element
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r_elem.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_shading(paragraph, color_hex):
    """Add background shading to a paragraph."""
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    paragraph.paragraph_format.element.get_or_add_pPr().append(shading)


def main():
    # Usage: python convert_to_docx.py [input.html] [output.docx]
    import sys
    input_file = sys.argv[1] if len(sys.argv) > 1 else "resume.html"
    output_file = sys.argv[2] if len(sys.argv) > 2 else "简历 钱佳宏.docx"
    print(f"Input:  {input_file}")
    print(f"Output: {output_file}")

    with open(input_file, "r", encoding="utf-8") as f:
        html = f.read()

    soup = BeautifulSoup(html, "html.parser")
    body = soup.find("body")
    doc = Document()

    # A4 page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    # Normal style
    style = doc.styles['Normal']
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    style.font.size = Pt(9.5)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.45

    # Tab stop for right-aligned dates
    tab_stops = style.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(14.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # ===== Process header =====
    header = body.find("div", class_="header")
    if header:
        # Title
        h1 = header.find("h1")
        if h1:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(h1.get_text(strip=True))
            set_font(run, size=22, bold=True, color=DARK)

        # Info table
        info_table = header.find("table", class_="info-table")
        if info_table:
            for row in info_table.find_all("tr"):
                cells = row.find_all("td")
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                parts = [c.get_text(strip=True) for c in cells if c.get_text(strip=True)]
                if parts:
                    run = p.add_run("  ".join(parts))
                    set_font(run, size=9, color=TEXT)

        # Links - table layout (4 cols: label, value, label, value)
        links_table = header.find("table", class_="links-table")
        if links_table:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            for row in links_table.find_all("tr"):
                cells = row.find_all("td")
                texts = []
                for i in range(0, len(cells), 2):
                    label = cells[i].get_text(strip=True)
                    a_tag = cells[i + 1].find("a") if i + 1 < len(cells) else None
                    if a_tag:
                        texts.append(f"{label}{a_tag.get_text(strip=True)}")
                    else:
                        texts.append(f"{label}{cells[i + 1].get_text(strip=True)}")
                run = p.add_run("    ".join(texts))
                set_font(run, size=8, color=LIGHT_GRAY)

    # ===== Process body children =====
    for child in body.children:
        if not hasattr(child, 'name') or child.name is None:
            continue

        classes = child.get("class", [])

        # Divider (skip)
        if child.name == "hr":
            continue

        # Section title
        if "sec-title" in classes:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"  {child.get_text(strip=True)}")
            set_font(run, size=11, bold=True, color=(255, 255, 255))
            add_shading(p, "2c3e50")

        # Entry (实习/项目经历)
        elif "entry" in classes:
            entry_top = child.find("div", class_="entry-top")
            if entry_top:
                name = entry_top.find("span", class_="entry-name")
                date = entry_top.find("span", class_="entry-date")
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(1)
                if name:
                    run = p.add_run(name.get_text(strip=True))
                    set_font(run, size=10, bold=True, color=DARK)
                if date:
                    run = p.add_run("\t")
                    set_font(run, size=8.5)
                    run = p.add_run(date.get_text(strip=True))
                    set_font(run, size=8.5, color=GRAY)

            role = child.find("div", class_="entry-role")
            if role:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(role.get_text(strip=True))
                set_font(run, size=9, bold=True, color=(52, 73, 94))

            text_div = child.find("div", class_="entry-text")
            if text_div:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.first_line_indent = Cm(0.75)
                run = p.add_run(text_div.get_text(strip=True))
                set_font(run, size=9.5, color=TEXT)

        # Skill row
        elif "skill" in classes:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            b_tag = child.find("b")
            if b_tag:
                run = p.add_run(b_tag.get_text(strip=True))
                set_font(run, size=9.5, bold=True, color=DARK)
                rest = child.get_text().replace(b_tag.get_text(), "", 1)
                run = p.add_run(rest)
                set_font(run, size=9.5, color=TEXT)
            else:
                run = p.add_run(child.get_text(strip=True))
                set_font(run, size=9.5, color=TEXT)

        # PR list
        elif "pr-list" in classes:
            for li in child.find_all("li"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(0.5)

                tag = li.find("span", class_="pr-tag")
                a_tag = li.find("a")
                full_text = li.get_text(strip=True)

                if tag:
                    run = p.add_run(f"[{tag.get_text(strip=True)}] ")
                    set_font(run, size=7.5, bold=True, color=(255, 255, 255))
                    add_shading(p, "2c3e50")

                if a_tag:
                    run = p.add_run(a_tag.get_text(strip=True))
                    set_font(run, size=9, bold=True, color=LINK)
                    desc = full_text.replace(a_tag.get_text(strip=True), "", 1).strip()
                    if desc.startswith("-"):
                        desc = desc[1:].strip()
                    run = p.add_run(f" - {desc}")
                    set_font(run, size=9, color=TEXT)

        # Self intro
        elif "intro" in classes:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0.75)
            run = p.add_run(child.get_text(strip=True))
            set_font(run, size=9.5, color=TEXT)

        # Footer
        elif "footer" in classes:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(child.get_text(strip=True))
            set_font(run, size=7.5, color=FOOTER_COLOR)

    doc.save(output_file)
    print("DOCX generated successfully.")


if __name__ == "__main__":
    main()
